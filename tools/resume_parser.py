"""
Resume parser tool for extracting structure and content from PDF/DOCX files.

This is a pure document processing tool that implements deterministic, stateless parsing
of resume files. It uses pdfplumber for PDF parsing and python-docx for DOCX parsing
to extract structured data for Resume objects.

This tool is designed to be used by the ResumeParserAgent for orchestration
within the LangGraph workflow.
"""

import re
from typing import List, Tuple, Dict, Any, Union
from pathlib import Path
from io import BytesIO

import pdfplumber
from docx import Document
from docx.document import Document as DocumentType

from src.schemas.core import Resume, ResumeBullet, ResumeSection


class ResumeParsingError(Exception):
    """Raised when resume parsing encounters unrecoverable errors."""

    pass


class ResumeParser:
    """
    Pure tool for parsing resume files into structured data.

    This tool provides stateless, deterministic parsing of PDF and DOCX resume files
    using document processing libraries. It extracts structure, bullets, skills,
    dates, and maintains span offsets for evidence tracking.

    Features:
    - PDF parsing using pdfplumber with robust text extraction
    - DOCX parsing using python-docx with style-aware extraction
    - Section detection using multiple heuristics (headers, formatting, patterns)
    - Bullet point extraction with original text span preservation
    - Skills extraction using pattern matching and keyword detection
    - Date extraction with multiple format support
    - Configurable filtering and validation
    """

    def __init__(self):
        """Initialize the parser with section detection patterns."""
        # Common section headers and patterns
        self.section_patterns = [
            r"^(experience|work experience|employment|professional experience)\s*:?\s*$",
            r"^(education|academic background|qualifications)\s*:?\s*$",
            r"^(skills|technical skills|core competencies|expertise)\s*:?\s*$",
            r"^(projects|key projects|notable projects)\s*:?\s*$",
            r"^(certifications|certificates|licenses)\s*:?\s*$",
            r"^(awards|achievements|honors|recognition)\s*:?\s*$",
            r"^(summary|objective|profile|about|overview)\s*:?\s*$",
            r"^(contact|contact information)\s*:?\s*$",
            r"^(publications|research|papers)\s*:?\s*$",
            r"^(volunteer|volunteering|community service)\s*:?\s*$",
        ]

        # Bullet point indicators
        self.bullet_patterns = [
            r"^\s*[•·▪▫◦‣⁃]\s*",  # Unicode bullets
            r"^\s*[-*+]\s*",  # Ascii bullets
            r"^\s*\d+[\.\)]\s*",  # Numbered lists
            r"^\s*[a-zA-Z][\.\)]\s*",  # Lettered lists
        ]

        # Date patterns (various formats)
        self.date_patterns = [
            r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\b",
            r"\b\d{1,2}/\d{1,2}/\d{2,4}\b",
            r"\b\d{4}\s*[-–—]\s*\d{4}\b",
            r"\b\d{4}\s*[-–—]\s*(?:present|current)\b",
            r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b",
            r"\b\d{4}\b",  # Simple 4-digit years
        ]

        # Skills indicators and tech terms
        self.skills_indicators = [
            "python",
            "java",
            "javascript",
            "react",
            "node",
            "sql",
            "aws",
            "docker",
            "kubernetes",
            "git",
            "linux",
            "machine learning",
            "data analysis",
            "project management",
            "agile",
            "scrum",
            "leadership",
            "communication",
        ]

    def parse_file(self, file_path: Union[str, Path]) -> Tuple[Resume, Dict[str, Any]]:
        """
        Parse a resume file and return structured Resume object with parsing report.

        Args:
            file_path: Path to PDF or DOCX file

        Returns:
            Tuple of (Resume object, parsing report dict)

        Raises:
            ResumeParsingError: If file cannot be parsed or is invalid format
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise ResumeParsingError(f"File not found: {file_path}")

        # Determine file type and parse accordingly
        file_ext = file_path.suffix.lower()

        try:
            if file_ext == ".pdf":
                raw_text = self._parse_pdf(file_path)
            elif file_ext in (".docx", ".doc"):
                raw_text = self._parse_docx(file_path)
            else:
                raise ResumeParsingError(f"Unsupported file format: {file_ext}")

            # Parse structure from raw text
            return self._parse_structure(raw_text)

        except Exception as e:
            raise ResumeParsingError(f"Error parsing {file_path}: {str(e)}") from e

    def parse_bytes(
        self, file_bytes: bytes, file_ext: str
    ) -> Tuple[Resume, Dict[str, Any]]:
        """
        Parse resume from bytes data.

        Args:
            file_bytes: Raw file bytes
            file_ext: File extension (.pdf, .docx, etc.)

        Returns:
            Tuple of (Resume object, parsing report dict)
        """
        file_ext = file_ext.lower()

        try:
            if file_ext == ".pdf":
                raw_text = self._parse_pdf_bytes(file_bytes)
            elif file_ext in (".docx", ".doc"):
                raw_text = self._parse_docx_bytes(file_bytes)
            else:
                raise ResumeParsingError(f"Unsupported file format: {file_ext}")

            return self._parse_structure(raw_text)

        except Exception as e:
            raise ResumeParsingError(f"Error parsing bytes: {str(e)}") from e

    def _parse_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file using pdfplumber."""
        text_parts = []

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text.strip())

        if not text_parts:
            raise ResumeParsingError("No text could be extracted from PDF")

        return "\n".join(text_parts)

    def _parse_pdf_bytes(self, file_bytes: bytes) -> str:
        """Extract text from PDF bytes using pdfplumber."""
        text_parts = []

        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text.strip())

        if not text_parts:
            raise ResumeParsingError("No text could be extracted from PDF")

        return "\n".join(text_parts)

    def _parse_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file using python-docx."""
        doc = Document(file_path)
        return self._extract_docx_text(doc)

    def _parse_docx_bytes(self, file_bytes: bytes) -> str:
        """Extract text from DOCX bytes using python-docx."""
        doc = Document(BytesIO(file_bytes))
        return self._extract_docx_text(doc)

    def _extract_docx_text(self, doc: DocumentType) -> str:
        """Extract structured text from DOCX document."""
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        if not text_parts:
            raise ResumeParsingError("No text could be extracted from DOCX")

        return "\n".join(text_parts)

    def _parse_structure(self, raw_text: str) -> Tuple[Resume, Dict[str, Any]]:
        """
        Parse raw text into structured Resume components.

        Args:
            raw_text: Raw extracted text from document

        Returns:
            Tuple of (Resume object, parsing report)
        """
        if not raw_text.strip():
            raise ResumeParsingError("Empty document")

        # Extract components
        sections = self._extract_sections(raw_text)
        bullets = self._extract_bullets(raw_text, sections)
        skills = self._extract_skills(raw_text)
        dates = self._extract_dates(raw_text)

        # Create Resume object
        resume = Resume(
            raw_text=raw_text,
            bullets=bullets,
            skills=skills,
            dates=dates,
            sections=sections,
        )

        # Generate parsing report
        report = {
            "text_length": len(raw_text),
            "sections_found": len(sections),
            "bullets_found": len(bullets),
            "skills_found": len(skills),
            "dates_found": len(dates),
            "confidence": self._calculate_confidence(resume),
            "warnings": self._generate_warnings(resume),
        }

        return resume, report

    def _extract_sections(self, text: str) -> List[ResumeSection]:
        """Extract resume sections with boundaries."""
        sections = []
        lines = text.split("\n")
        section_starts = []  # (line_index, section_name)

        # Find all section headers
        for i, line in enumerate(lines):
            line_stripped = line.strip().lower()

            # Check if line matches any section pattern
            is_section_header = any(
                re.match(pattern, line_stripped, re.IGNORECASE)
                for pattern in self.section_patterns
            )

            # Also check for short lines that look like headers (all caps, etc.)
            if not is_section_header and len(line_stripped) < 30:
                if line.strip().isupper() or (
                    len(line_stripped) > 3 and line_stripped.replace(" ", "").isalpha()
                ):
                    is_section_header = True

            if is_section_header:
                section_starts.append((i, line.strip()))

        # Create sections with proper boundaries
        for i, (line_idx, section_name) in enumerate(section_starts):
            # Calculate character offset to start of this section header line
            start_offset = sum(len(lines[j]) + 1 for j in range(line_idx))

            # Calculate end offset (start of next section or end of text)
            if i + 1 < len(section_starts):
                next_line_idx = section_starts[i + 1][0]
                end_offset = sum(len(lines[j]) + 1 for j in range(next_line_idx))
            else:
                end_offset = len(text)

            sections.append(
                ResumeSection(
                    name=section_name,
                    bullets=[],  # Will be filled later
                    start_offset=start_offset,
                    end_offset=end_offset,
                )
            )

        # If no sections found, create a default one
        if not sections:
            sections.append(
                ResumeSection(
                    name="Content", bullets=[], start_offset=0, end_offset=len(text)
                )
            )

        return sections

    def _extract_bullets(
        self, text: str, sections: List[ResumeSection]
    ) -> List[ResumeBullet]:
        """Extract bullet points with section assignment and span tracking."""
        bullets = []
        lines = text.split("\n")
        char_offset = 0

        for line in lines:
            # Check if line is a bullet point
            is_bullet = any(re.match(pattern, line) for pattern in self.bullet_patterns)

            if is_bullet:
                # Clean bullet text
                bullet_text = line.strip()
                for pattern in self.bullet_patterns:
                    bullet_text = re.sub(pattern, "", bullet_text).strip()

                if bullet_text:  # Only add non-empty bullets
                    # Find which section this bullet belongs to
                    section_name = self._find_bullet_section(char_offset, sections)

                    bullet = ResumeBullet(
                        text=bullet_text,
                        section=section_name,
                        start_offset=char_offset,
                        end_offset=char_offset + len(line),
                    )
                    bullets.append(bullet)

                    # Add bullet to its section
                    for section in sections:
                        if section.name == section_name:
                            section.bullets.append(bullet)
                            break

            char_offset += len(line) + 1  # +1 for newline

        return bullets

    def _find_bullet_section(
        self, char_offset: int, sections: List[ResumeSection]
    ) -> str:
        """Find which section a bullet belongs to based on character offset."""
        for section in sections:
            if section.start_offset <= char_offset <= section.end_offset:
                return section.name

        # Fallback to first section or default
        return sections[0].name if sections else "Content"

    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from text using pattern matching and keyword detection."""
        skills = set()
        text_lower = text.lower()

        # Look for explicit skills sections
        lines = text.split("\n")
        in_skills_section = False

        for line in lines:
            line_stripped = line.strip().lower()

            # Check if entering skills section
            if any(
                "skill" in line_stripped and pattern in line_stripped
                for pattern in [":", "skill", "competenc", "expertis"]
            ):
                in_skills_section = True
                continue

            # Check if leaving skills section (new section header)
            if in_skills_section and len(line_stripped) < 30 and line.strip().isupper():
                in_skills_section = False
                continue

            # Extract skills from skills section
            if in_skills_section:
                # Split on common delimiters
                line_skills = re.split(r"[,;|•·▪▫◦‣⁃]+", line_stripped)
                for skill in line_skills:
                    skill = skill.strip()
                    if len(skill) > 1 and skill not in ["and", "or", "with"]:
                        skills.add(skill.title())

        # Also look for known technical skills throughout text
        for indicator in self.skills_indicators:
            if indicator.lower() in text_lower:
                skills.add(indicator.title())

        return sorted(list(skills))

    def _extract_dates(self, text: str) -> List[str]:
        """Extract dates from text using multiple pattern matching."""
        dates = set()

        for pattern in self.date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                dates.add(match.strip())

        return sorted(list(dates))

    def _calculate_confidence(self, resume: Resume) -> float:
        """Calculate parsing confidence based on extracted content."""
        score = 0.0

        # Base score for successful parsing
        if resume.raw_text:
            score += 0.3

        # Bonus for sections found
        if len(resume.sections) > 1:
            score += 0.2
        if len(resume.sections) > 3:
            score += 0.1

        # Bonus for bullets found
        if len(resume.bullets) > 0:
            score += 0.2
        if len(resume.bullets) > 5:
            score += 0.1

        # Bonus for skills found
        if len(resume.skills) > 0:
            score += 0.1

        # Bonus for dates found
        if len(resume.dates) > 0:
            score += 0.1

        return min(score, 1.0)

    def _generate_warnings(self, resume: Resume) -> List[str]:
        """Generate warnings about parsing quality."""
        warnings = []

        if len(resume.sections) == 0:
            warnings.append("No resume sections detected")
        elif len(resume.sections) == 1:
            warnings.append(
                "Only one section detected - resume structure may be unclear"
            )

        if len(resume.bullets) == 0:
            warnings.append("No bullet points detected")
        elif len(resume.bullets) < 3:
            warnings.append("Few bullet points found - resume may be sparse")

        if len(resume.skills) == 0:
            warnings.append("No skills detected")

        if len(resume.dates) == 0:
            warnings.append("No dates found - experience timeline unclear")

        if len(resume.raw_text) < 500:
            warnings.append("Resume text is quite short")
        elif len(resume.raw_text) > 10000:
            warnings.append("Resume text is quite long (>2 pages)")

        return warnings
